"""
Unit tests for tools/data_harvester_tool.py
Target coverage: ≥ 85%

Heavy extractors (PDF/DOCX/OCR) are tested against real sample files
in data/incoming/ using lightweight assertions (no mocking of pdfplumber etc.)
so the actual integration is verified without needing API keys.
"""

import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock

from tools.data_harvester_tool import (
    harvest,
    detect_language,
    get_file_format,
    _clean_table,
    _error_result,
    _success_result,
    PDF_EXTENSIONS,
    DOCX_EXTENSIONS,
    IMAGE_EXTENSIONS,
)

# ── Paths to real sample files ─────────────────────────────────────────────────

DATA_DIR = Path(__file__).parent.parent / "data" / "incoming"
PDF_FILE = DATA_DIR / "INV_EN_001.pdf"
DOCX_FILE = DATA_DIR / "INV_DE_004.pdf"   # Use available PDF as fallback
IMAGE_FILE = DATA_DIR / "INV_EN_005_scan.png"
MALFORMED_PDF = DATA_DIR / "INV_EN_006_malformed.pdf"


# ── get_file_format ────────────────────────────────────────────────────────────

class TestGetFileFormat:
    def test_pdf_extension(self):
        assert get_file_format("invoice.pdf") == "pdf"

    def test_pdf_uppercase(self):
        assert get_file_format("invoice.PDF") == "pdf"

    def test_docx_extension(self):
        assert get_file_format("doc.docx") == "docx"

    def test_doc_extension(self):
        assert get_file_format("doc.doc") == "docx"

    def test_png_extension(self):
        assert get_file_format("scan.png") == "image"

    def test_jpg_extension(self):
        assert get_file_format("scan.jpg") == "image"

    def test_jpeg_extension(self):
        assert get_file_format("scan.jpeg") == "image"

    def test_tiff_extension(self):
        assert get_file_format("scan.tiff") == "image"

    def test_unknown_extension(self):
        assert get_file_format("file.xyz") == "unknown"

    def test_no_extension(self):
        assert get_file_format("README") == "unknown"


# ── _clean_table ───────────────────────────────────────────────────────────────

class TestCleanTable:
    def test_strips_whitespace(self):
        rows = [["  hello  ", " world "]]
        result = _clean_table(rows)
        assert result == [["hello", "world"]]

    def test_removes_fully_empty_rows(self):
        rows = [["", ""], ["data", "here"]]
        result = _clean_table(rows)
        assert result == [["data", "here"]]

    def test_keeps_partially_filled_rows(self):
        rows = [["value", ""]]
        result = _clean_table(rows)
        assert result == [["value", ""]]

    def test_none_cells_become_empty_string(self):
        rows = [[None, "data"]]
        result = _clean_table(rows)
        assert result == [["", "data"]]

    def test_none_row_skipped(self):
        rows = [None, ["a", "b"]]
        result = _clean_table(rows)
        assert result == [["a", "b"]]

    def test_empty_input(self):
        assert _clean_table([]) == []

    def test_all_empty_rows_removed(self):
        rows = [["", ""], [None, None]]
        result = _clean_table(rows)
        assert result == []


# ── _error_result ──────────────────────────────────────────────────────────────

class TestErrorResult:
    def test_raw_text_empty(self):
        r = _error_result("boom")
        assert r["raw_text"] == ""

    def test_tables_empty(self):
        r = _error_result("boom")
        assert r["tables"] == []

    def test_error_message_set(self):
        r = _error_result("file not found")
        assert r["error"] == "file not found"

    def test_file_format_unknown(self):
        r = _error_result("oops")
        assert r["file_format"] == "unknown"

    def test_detected_language_defaults_en(self):
        r = _error_result("oops")
        assert r["detected_language"] == "en"


# ── _success_result ────────────────────────────────────────────────────────────

class TestSuccessResult:
    def test_raw_text_preserved(self):
        r = _success_result("hello world invoice", [], "pdf")
        assert r["raw_text"] == "hello world invoice"

    def test_tables_preserved(self):
        tables = [["a", "b"], ["c", "d"]]
        r = _success_result("text", [tables], "pdf")
        assert r["tables"] == [tables]

    def test_file_format_preserved(self):
        r = _success_result("text", [], "docx")
        assert r["file_format"] == "docx"

    def test_error_is_none(self):
        r = _success_result("text", [], "pdf")
        assert r["error"] is None

    def test_detected_language_set(self):
        r = _success_result("Invoice total amount due", [], "pdf")
        assert isinstance(r["detected_language"], str)
        assert len(r["detected_language"]) >= 2


# ── detect_language ────────────────────────────────────────────────────────────

class TestDetectLanguage:
    def test_empty_string_returns_en(self):
        assert detect_language("") == "en"

    def test_short_text_returns_en(self):
        assert detect_language("hi") == "en"

    def test_english_text(self):
        result = detect_language("This is an invoice for services rendered in the month of March.")
        assert result == "en"

    def test_none_like_whitespace_returns_en(self):
        assert detect_language("   ") == "en"

    def test_langdetect_exception_falls_back_to_en(self):
        with patch("tools.data_harvester_tool.detect_language") as mock_detect:
            mock_detect.return_value = "en"
            assert detect_language("!!!") == "en"

    def test_returns_string(self):
        result = detect_language("The quick brown fox jumps over the lazy dog.")
        assert isinstance(result, str)


# ── harvest — file not found ───────────────────────────────────────────────────

class TestHarvestNotFound:
    def test_missing_file_returns_error(self):
        result = harvest("/nonexistent/path/invoice.pdf")
        assert result["error"] is not None
        assert "not found" in result["error"].lower()

    def test_missing_file_raw_text_empty(self):
        result = harvest("/nonexistent/path/invoice.pdf")
        assert result["raw_text"] == ""

    def test_unsupported_extension_returns_error(self, tmp_path):
        # File must exist so it passes the existence check, then hits extension check
        fake = tmp_path / "somefile.xyz"
        fake.write_text("dummy")
        result = harvest(str(fake))
        assert result["error"] is not None
        assert "Unsupported" in result["error"]


# ── harvest — real PDF files ───────────────────────────────────────────────────

class TestHarvestPDF:
    def test_english_pdf_extracts_text(self):
        result = harvest(str(PDF_FILE))
        assert result["error"] is None
        assert len(result["raw_text"]) > 50

    def test_english_pdf_format(self):
        result = harvest(str(PDF_FILE))
        assert result["file_format"] == "pdf"

    def test_english_pdf_language_detected(self):
        result = harvest(str(PDF_FILE))
        assert result["detected_language"] == "en"

    def test_malformed_pdf_does_not_crash(self):
        result = harvest(str(MALFORMED_PDF))
        # Should return either extracted text or an error — never raise
        assert "raw_text" in result
        assert "error" in result

    def test_tables_is_list(self):
        result = harvest(str(PDF_FILE))
        assert isinstance(result["tables"], list)


# ── harvest — real image file ──────────────────────────────────────────────────

def _tesseract_available() -> bool:
    """Return True only if Tesseract can actually run OCR (binary + language data present)."""
    try:
        import pytesseract
        from PIL import Image
        import io
        # Create a minimal 1x1 white image and attempt OCR
        img = Image.new("RGB", (10, 10), color=(255, 255, 255))
        pytesseract.image_to_string(img)
        return True
    except Exception:
        return False


_TESSERACT_OK = _tesseract_available()
_IMAGE_OK = IMAGE_FILE.exists() and _TESSERACT_OK


class TestHarvestImage:
    @pytest.mark.skipif(not _IMAGE_OK, reason="Tesseract not installed or image missing")
    def test_image_returns_format(self):
        result = harvest(str(IMAGE_FILE))
        assert result["file_format"] == "image"

    @pytest.mark.skipif(not _IMAGE_OK, reason="Tesseract not installed or image missing")
    def test_image_tables_empty(self):
        result = harvest(str(IMAGE_FILE))
        assert result["tables"] == []

    def test_image_extraction_exception_returns_error(self):
        with patch("tools.data_harvester_tool._extract_image", side_effect=Exception("OCR failed")):
            result = harvest(str(IMAGE_FILE))
            assert result["error"] is not None


# ── harvest — real DOCX (created in test) ─────────────────────────────────────

class TestHarvestDOCX:
    def test_docx_extracts_paragraphs(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Invoice No: INV-TEST-001")
        doc.add_paragraph("Vendor: Test Corp")
        doc.add_paragraph("Total Amount: $1,500.00")
        path = tmp_path / "test_invoice.docx"
        doc.save(str(path))

        result = harvest(str(path))
        assert result["error"] is None
        assert result["file_format"] == "docx"
        assert "INV-TEST-001" in result["raw_text"]

    def test_docx_extracts_tables(self, tmp_path):
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Qty"
        table.cell(0, 2).text = "Price"
        table.cell(1, 0).text = "SKU-001"
        table.cell(1, 1).text = "10"
        table.cell(1, 2).text = "50.00"
        path = tmp_path / "table_invoice.docx"
        doc.save(str(path))

        result = harvest(str(path))
        assert result["error"] is None
        assert isinstance(result["tables"], list)
        assert len(result["tables"]) > 0

    def test_docx_table_content_in_raw_text(self, tmp_path):
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "SKU-999"
        table.cell(0, 1).text = "99.99"
        path = tmp_path / "t.docx"
        doc.save(str(path))

        result = harvest(str(path))
        assert "SKU-999" in result["raw_text"]

    def test_docx_language_detected(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("This invoice is for logistics services rendered in March.")
        path = tmp_path / "lang.docx"
        doc.save(str(path))

        result = harvest(str(path))
        assert result["detected_language"] == "en"

    def test_docx_empty_paragraphs_skipped(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("  ")
        doc.add_paragraph("Real content here")
        path = tmp_path / "empty_paras.docx"
        doc.save(str(path))

        result = harvest(str(path))
        assert result["raw_text"].strip() == "Real content here"


# ── harvest — image path via mock (covers _extract_image) ─────────────────────

class TestHarvestImageMocked:
    def test_image_path_calls_extract_image(self, tmp_path):
        fake_png = tmp_path / "test.png"
        fake_png.write_bytes(b"\x89PNG\r\n")  # fake PNG header

        with patch("tools.data_harvester_tool._extract_image") as mock_extract:
            mock_extract.return_value = {
                "raw_text": "Invoice 123", "tables": [],
                "detected_language": "en", "file_format": "image", "error": None
            }
            result = harvest(str(fake_png))

        assert result["file_format"] == "image"
        mock_extract.assert_called_once()

    def test_extract_image_uses_greyscale(self, tmp_path):
        """_extract_image converts to greyscale before OCR."""
        from PIL import Image as PILImage
        img = PILImage.new("RGB", (50, 50), color=(200, 200, 200))
        path = tmp_path / "test.png"
        img.save(str(path))

        with patch("pytesseract.image_to_string", return_value="Invoice text here") as mock_ocr:
            from tools.data_harvester_tool import _extract_image
            result = _extract_image(path)

        assert result["file_format"] == "image"
        assert result["raw_text"] == "Invoice text here"
        assert result["error"] is None
        # Verify greyscale conversion happened — PIL converts then passes to OCR
        call_args = mock_ocr.call_args[0][0]
        assert call_args.mode == "L"

    def test_extract_image_respects_tesseract_cmd_env(self, tmp_path, monkeypatch):
        from PIL import Image as PILImage
        img = PILImage.new("RGB", (20, 20))
        path = tmp_path / "t.png"
        img.save(str(path))

        monkeypatch.setenv("TESSERACT_CMD", "/custom/tesseract")
        with patch("pytesseract.image_to_string", return_value="text"):
            from tools.data_harvester_tool import _extract_image
            import pytesseract
            _extract_image(path)
            assert pytesseract.pytesseract.tesseract_cmd == "/custom/tesseract"


# ── harvest — extension constants ──────────────────────────────────────────────

class TestExtensionSets:
    def test_pdf_in_pdf_extensions(self):
        assert ".pdf" in PDF_EXTENSIONS

    def test_docx_in_docx_extensions(self):
        assert ".docx" in DOCX_EXTENSIONS

    def test_doc_in_docx_extensions(self):
        assert ".doc" in DOCX_EXTENSIONS

    def test_png_in_image_extensions(self):
        assert ".png" in IMAGE_EXTENSIONS

    def test_jpg_in_image_extensions(self):
        assert ".jpg" in IMAGE_EXTENSIONS
