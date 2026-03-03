"""
Data Harvester Tool — AI Invoice Auditor
Extracts raw text and tables from invoice files in three formats:
  - PDF  → pdfplumber
  - DOCX → python-docx
  - Image (PNG/JPG) → pytesseract OCR

Also detects the source language of the extracted text.
"""

import os
import re
from pathlib import Path
from typing import Optional

from core.logger import get_logger

logger = get_logger(__name__)

# Supported file extensions grouped by extraction method
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx", ".doc"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}


# ── Public API ─────────────────────────────────────────────────────────────────

def harvest(file_path: str) -> dict:
    """
    Extract text from an invoice file and detect its language.

    Args:
        file_path: Absolute or relative path to the invoice file.

    Returns:
        {
            "raw_text":          str,   # full extracted text
            "tables":            list,  # list of table rows (list[list[str]])
            "detected_language": str,   # ISO 639-1 code e.g. "en", "es", "de"
            "file_format":       str,   # "pdf" | "docx" | "image"
            "error":             str | None
        }
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if not path.exists():
        return _error_result(f"File not found: {file_path}")

    try:
        if ext in PDF_EXTENSIONS:
            return _extract_pdf(path)
        elif ext in DOCX_EXTENSIONS:
            return _extract_docx(path)
        elif ext in IMAGE_EXTENSIONS:
            return _extract_image(path)
        else:
            return _error_result(f"Unsupported file type: {ext}")
    except Exception as exc:
        logger.error("Extraction failed for %s: %s", file_path, exc)
        return _error_result(str(exc))


def detect_language(text: str) -> str:
    """
    Detect the language of a text string.
    Returns ISO 639-1 code (e.g. 'en', 'es', 'de').
    Falls back to 'en' if detection fails or text is too short.
    """
    if not text or len(text.strip()) < 20:
        return "en"
    try:
        from langdetect import detect, LangDetectException
        return detect(text)
    except Exception:
        return "en"


def get_file_format(file_path: str) -> str:
    """Return normalised format string for a file path."""
    ext = Path(file_path).suffix.lower()
    if ext in PDF_EXTENSIONS:
        return "pdf"
    if ext in DOCX_EXTENSIONS:
        return "docx"
    if ext in IMAGE_EXTENSIONS:
        return "image"
    return "unknown"


# ── Extractors ─────────────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> dict:
    import pdfplumber

    text_parts: list[str] = []
    tables: list[list[list[str]]] = []

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)

            for table in page.extract_tables() or []:
                cleaned = _clean_table(table)
                if cleaned:
                    tables.append(cleaned)

    raw_text = "\n".join(text_parts).strip()
    logger.info("PDF extracted: %s (%d chars, %d tables)", path.name, len(raw_text), len(tables))
    return _success_result(raw_text, tables, "pdf")


def _extract_docx(path: Path) -> dict:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    text_parts: list[str] = []
    tables: list[list[list[str]]] = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        cleaned = _clean_table(rows)
        if cleaned:
            tables.append(cleaned)
            # Also include table text in raw_text
            for row in cleaned:
                text_parts.append(" | ".join(row))

    raw_text = "\n".join(text_parts).strip()
    logger.info("DOCX extracted: %s (%d chars, %d tables)", path.name, len(raw_text), len(tables))
    return _success_result(raw_text, tables, "docx")


def _extract_image(path: Path) -> dict:
    import pytesseract
    from PIL import Image

    # Use TESSERACT_CMD env var if set (Windows)
    tesseract_cmd = os.environ.get("TESSERACT_CMD")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    img = Image.open(str(path))
    # Pre-process: convert to greyscale for better OCR accuracy
    img = img.convert("L")

    try:
        raw_text = pytesseract.image_to_string(img, config="--psm 6").strip()
        if raw_text:
            logger.info("Image OCR extracted: %s (%d chars)", path.name, len(raw_text))
            return _success_result(raw_text, [], "image")
    except Exception as exc:
        logger.warning("Tesseract OCR failed for %s: %s", path.name, exc)

    # Fallback to EasyOCR if Tesseract is not available
    try:
        import easyocr
        reader = easyocr.Reader(["en"], gpu=False)
        lines = reader.readtext(str(path), detail=0)
        raw_text = "\n".join(lines).strip()
        logger.info("EasyOCR extracted: %s (%d chars)", path.name, len(raw_text))
        return _success_result(raw_text, [], "image")
    except Exception as exc:
        logger.error("EasyOCR failed for %s: %s", path.name, exc)
        return _error_result(f"OCR failed: {exc}")


# ── Helpers ────────────────────────────────────────────────────────────────────

def _clean_table(rows: list) -> list[list[str]]:
    """Remove fully-empty rows and normalise cell whitespace."""
    cleaned = []
    for row in rows:
        if row is None:
            continue
        cells = [str(cell).strip() if cell is not None else "" for cell in row]
        if any(cells):
            cleaned.append(cells)
    return cleaned


def _success_result(raw_text: str, tables: list, file_format: str) -> dict:
    lang = detect_language(raw_text)
    return {
        "raw_text": raw_text,
        "tables": tables,
        "detected_language": lang,
        "file_format": file_format,
        "error": None,
    }


def _error_result(message: str) -> dict:
    return {
        "raw_text": "",
        "tables": [],
        "detected_language": "en",
        "file_format": "unknown",
        "error": message,
    }
